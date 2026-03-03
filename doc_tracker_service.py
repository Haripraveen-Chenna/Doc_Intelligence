from doc_initialize import get_preset_list, agents
import os
import re
import hashlib
import asyncio
import logging
from functools import wraps
from typing import Optional

import numpy as np
import pandas as pd
from io import StringIO, BytesIO

from langchain_anthropic import ChatAnthropic
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate
from dotenv import load_dotenv
from docx import Document

from fastapi import UploadFile
from pydantic import BaseModel, field_validator
from dateutil import parser as dateutil_parser
from tenacity import (
    retry,
    stop_after_attempt,
    wait_exponential,
    retry_if_exception_type,
    before_sleep_log,
)
import anthropic  # for specific exception types

load_dotenv()
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
logger = logging.getLogger(__name__)

_parser = JsonOutputParser()
CHUNK_SIZE = 15

# ============================================================
# FIX 4: LLM CLIENT FACTORY (no global mutable state)
# ============================================================
# Returns a fresh ChatAnthropic instance per call.
# This eliminates shared state across concurrent async tasks,
# which can cause subtle race conditions under heavy load.

def _make_haiku() -> ChatAnthropic:
    return ChatAnthropic(
        model_name="claude-haiku-4-5",
        api_key=ANTHROPIC_API_KEY,
    )

def _make_sonnet() -> ChatAnthropic:
    return ChatAnthropic(
        model_name="claude-sonnet-4-5",
        api_key=ANTHROPIC_API_KEY,
        max_tokens_to_sample=5000,
    )


# ============================================================
# FIX 3: RETRY / BACKOFF DECORATOR
# ============================================================
# Wraps any coroutine with:
#   - Up to 4 attempts
#   - Exponential backoff: 2s → 4s → 8s (max 30s)
#   - Only retries on Anthropic API / network errors
#   - Logs each retry attempt for observability

_RETRYABLE_EXCEPTIONS = (
    anthropic.APIConnectionError,
    anthropic.APITimeoutError,
    anthropic.RateLimitError,
    anthropic.InternalServerError,
)

def with_retry(coro_fn):
    """Decorator that adds exponential backoff retry to an async function."""
    @retry(
        reraise=True,
        stop=stop_after_attempt(4),
        wait=wait_exponential(multiplier=1, min=2, max=30),
        retry=retry_if_exception_type(_RETRYABLE_EXCEPTIONS),
        before_sleep=before_sleep_log(logger, logging.WARNING),
    )
    @wraps(coro_fn)
    async def wrapper(*args, **kwargs):
        return await coro_fn(*args, **kwargs)
    return wrapper


# ============================================================
# FIX 2: DOCUMENT CONTENT CACHE
# ============================================================
# Keyed by SHA-256 of raw file bytes. Avoids re-paying full
# LLM cost when the same document is submitted twice in the
# same server process lifetime.
#
# For multi-process / persistent caching, swap _DOC_CACHE
# for a Redis client using the same key scheme.

_DOC_CACHE: dict[str, dict] = {}

def _cache_key(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


# ============================================================
# BUILD VALID AGENT LOOKUP FROM SOURCE OF TRUTH
# ============================================================

VALID_AGENTS: dict[int, str] = {a["id"]: a["name"] for a in agents}


# ============================================================
# STRICT STATUS NORMALIZATION
# ============================================================

ALLOWED_STATUSES = {"completed", "in_progress", "not_started", "planned", "blocked", "unknown"}

STATUS_MAP = {
    "done": "completed",
    "complete": "completed",
    "completed": "completed",
    "in progress": "in_progress",
    "inprocess": "in_progress",
    "in motion": "in_progress",
    "in development": "in_progress",
    "not started": "not_started",
    "pending": "not_started",
    "tbd": "not_started",
    "planning": "planned",
    "planned": "planned",
    "blocked": "blocked",
}

def normalize_status(status) -> str:
    if status is None:
        return "unknown"
    cleaned = str(status).lower().strip()
    if cleaned in ("??", "n/a", "-", "", "none", "null"):
        return "unknown"
    mapped = STATUS_MAP.get(cleaned)
    if mapped:
        return mapped
    for key, val in STATUS_MAP.items():
        if key in cleaned:
            return val
    return "unknown"


# ============================================================
# ISO DATE NORMALIZATION + EXCEL SERIAL DATE SUPPORT
# ============================================================

_NON_DATE_PATTERNS = re.compile(
    r"(confirm|tbd|pending|check|depends|asap|soon|ongoing|note|see|per|"
    r"after|before|when|if|\?|!)",
    re.IGNORECASE
)

_EXCEL_SERIAL_MIN = 32874
_EXCEL_SERIAL_MAX = 73050

def normalize_date(raw) -> Optional[str]:
    if raw is None:
        return None
    raw_str = str(raw).strip()
    if not raw_str or raw_str.lower() in ("null", "none", "n/a", "-", "??", "tbd"):
        return None
    if _NON_DATE_PATTERNS.search(raw_str):
        return None
    if re.fullmatch(r"\d{5}", raw_str):
        serial = int(raw_str)
        if _EXCEL_SERIAL_MIN <= serial <= _EXCEL_SERIAL_MAX:
            try:
                return pd.to_datetime(float(serial), origin="1899-12-30", unit="D").strftime("%Y-%m-%d")
            except Exception:
                return None
    if re.fullmatch(r"\d{1,2}", raw_str):
        return None
    try:
        parsed = dateutil_parser.parse(raw_str, dayfirst=True, yearfirst=False)
        return parsed.strftime("%Y-%m-%d")
    except (ValueError, OverflowError):
        return None


# ============================================================
# BUDGET NORMALIZATION
# ============================================================

CURRENCY_SYMBOLS = {
    "₹": "INR", "inr": "INR",
    "$": "USD", "usd": "USD",
    "€": "EUR", "eur": "EUR",
    "£": "GBP", "gbp": "GBP",
}

_BUDGET_KEYWORD_PATTERN = re.compile(
    r"(budget|cost|price|amount|fee|charge|spend|expenditure|₹|\$|€|£|inr|usd|eur|gbp|lakh|lac\b)",
    re.IGNORECASE
)

_BUDGET_VALUE_PATTERN = re.compile(
    r"(₹|\$|€|£|inr|usd|eur|gbp)?\s*([\d,]+(?:\.\d+)?)\s*(₹|\$|€|£|inr|usd|eur|gbp|lakh|lac|k)?",
    re.IGNORECASE
)

def extract_budget(value) -> Optional[dict]:
    if value is None:
        return None
    raw = str(value).strip()
    if not raw or raw.lower() in ("null", "none", "n/a", "-", "??"):
        return None
    if not _BUDGET_KEYWORD_PATTERN.search(raw):
        return None
    match = _BUDGET_VALUE_PATTERN.search(raw)
    if not match:
        return None
    prefix_sym = (match.group(1) or "").lower().strip()
    amount_str = (match.group(2) or "").replace(",", "")
    suffix_sym = (match.group(3) or "").lower().strip()
    try:
        amount = float(amount_str)
    except ValueError:
        return None
    if amount <= 0:
        return None
    if "lakh" in suffix_sym or "lac" in suffix_sym:
        amount *= 100_000
    elif suffix_sym == "k":
        amount *= 1_000
    currency = "INR"
    for sym, code in CURRENCY_SYMBOLS.items():
        if sym in prefix_sym or sym in suffix_sym:
            currency = code
            break
    return {"estimated": amount, "currency": currency}


# ============================================================
# AGENT VALIDATION AGAINST SOURCE OF TRUTH
# ============================================================

def validate_agent(agent: Optional[dict]) -> Optional[dict]:
    if not agent:
        return None
    agent_id = agent.get("id")
    agent_name = agent.get("name")
    if agent_id is None or agent_name is None:
        return None
    try:
        agent_id = int(agent_id)
    except (ValueError, TypeError):
        return None
    if agent_id not in VALID_AGENTS:
        return None
    expected_name = VALID_AGENTS[agent_id].lower().strip()
    if agent_name.lower().strip() != expected_name:
        agent["name"] = VALID_AGENTS[agent_id]
        agent["id"] = agent_id
    return agent


# ============================================================
# ASSIGNED_TO - SPLIT MULTI-PERSON STRINGS INTO LIST
# ============================================================

_ASSIGNED_SPLIT_PATTERN = re.compile(r"\s*[,/&+]\s*|\s+and\s+", re.IGNORECASE)
_GARBAGE_ASSIGNED = {"??", "n/a", "-", "", "none", "null", "tbd"}

def normalize_assigned_to(value) -> list[str]:
    if value is None:
        return []
    
    # If the LLM already passed a clean list, process its elements.
    if isinstance(value, list):
        flat_list = []
        for item in value:
            # Recursively handle nested stringified lists, e.g., ["['Ritika'", "'Sam']"]
            cleaned = str(item).strip("[]'\" ")
            if cleaned and cleaned.lower() not in _GARBAGE_ASSIGNED:
                flat_list.extend(_ASSIGNED_SPLIT_PATTERN.split(cleaned))
        
        final = [p.strip("[]'\" ").strip() for p in flat_list]
        final = [p for p in final if p and p.lower() not in _GARBAGE_ASSIGNED]
        return final

    raw = str(value).strip("[]'\" ")
    if not raw or raw.lower() in _GARBAGE_ASSIGNED:
        return []
        
    parts = _ASSIGNED_SPLIT_PATTERN.split(raw)
    cleaned = [p.strip("[]'\" ").rstrip(",").strip() for p in parts]
    result = [p for p in cleaned if p and p.lower() not in _GARBAGE_ASSIGNED]
    return result


# ============================================================
# ORIGIN NORMALIZATION
# ============================================================

def normalize_origin(origin_value) -> dict:
    if str(origin_value) == "1":
        return {"type": "user_upload", "confidence": 1.0}
    return {"type": "ai_generated", "confidence": 0.8}


# ============================================================
# COMPOSITE DEDUPLICATION (DROPS DUPLICATES ENTIRELY)
# ============================================================

def _task_signature(task: dict) -> str:
    name = re.sub(r"\s+", " ", (task.get("name") or "").strip().lower())
    details = task.get("details") or {}
    assignees = details.get("assigned_to") or []
    assignee_str = ",".join(sorted(a.lower() for a in assignees)) if isinstance(assignees, list) else ""
    deadline = (details.get("deadline") or "").strip()
    composite = f"{name}|{assignee_str}|{deadline}"
    return hashlib.md5(composite.encode()).hexdigest()

def _token_similarity(a: str, b: str) -> float:
    tokens_a = set(re.sub(r"[^a-z0-9 ]", "", a.lower()).split())
    tokens_b = set(re.sub(r"[^a-z0-9 ]", "", b.lower()).split())
    if not tokens_a or not tokens_b:
        return 0.0
    intersection = len(tokens_a & tokens_b)
    smallest_set = min(len(tokens_a), len(tokens_b))
    return intersection / smallest_set if smallest_set > 0 else 0.0

SIMILARITY_THRESHOLD = 0.70

def deduplicate_tasks(tasks: list) -> list:
    seen_sigs: set[str] = set()
    seen_names: list[str] = []
    result = []
    for task in tasks:
        name = (task.get("name") or "").strip()
        if not name:
            continue
        sig = _task_signature(task)
        if sig in seen_sigs:
            continue
        is_duplicate = False
        for existing_name in seen_names:
            if _token_similarity(name, existing_name) >= SIMILARITY_THRESHOLD:
                is_duplicate = True
                break
        if is_duplicate:
            continue
        seen_sigs.add(sig)
        seen_names.append(name)
        result.append(task)
    return result


# ============================================================
# FULL TASK NORMALIZATION PIPELINE
# ============================================================

def normalize_task(task: dict) -> Optional[dict]:
    name = (task.get("name") or "").strip()
    if not name:
        return None
    task["origin"] = normalize_origin(task.get("origin"))
    task["agent"] = validate_agent(task.get("agent"))
    details = task.get("details") or {}
    details["status"] = normalize_status(details.get("status"))
    details["deadline"] = normalize_date(details.get("deadline"))
    details["assigned_to"] = normalize_assigned_to(details.get("assigned_to"))
    raw_budget = details.pop("budget", None) or task.pop("budget", None)
    details["budget"] = extract_budget(raw_budget)
    task["details"] = details
    return task

def _normalize_tasks(tasks: list) -> list:
    normalized = [normalize_task(t) for t in tasks]
    normalized = [t for t in normalized if t]
    return deduplicate_tasks(normalized)


# ============================================================
# Pydantic Models
# ============================================================

class AgentModel(BaseModel):
    name: str
    id: int

class BudgetModel(BaseModel):
    estimated: float
    currency: str = "INR"

class TaskDetails(BaseModel):
    assigned_to: Optional[list[str]] = None
    deadline: Optional[str] = None
    status: str = "unknown"
    budget: Optional[BudgetModel] = None

    @field_validator("status", mode="before")
    @classmethod
    def enforce_status(cls, v):
        return normalize_status(v)

    @field_validator("deadline", mode="before")
    @classmethod
    def enforce_date(cls, v):
        return normalize_date(v)

class OriginModel(BaseModel):
    type: str
    confidence: float

class TaskModel(BaseModel):
    name: str
    description: str = ""
    origin: OriginModel
    agent: Optional[AgentModel] = None
    details: Optional[TaskDetails] = None

class ProjectModel(BaseModel):
    project_name: str
    preset_name: str
    tasks: list[TaskModel]


# ============================================================
# File reading
# ============================================================

def detect_header_row(df: pd.DataFrame, min_non_null: int = 2) -> int:
    for i in range(min(15, len(df))):
        if df.iloc[i].notna().sum() >= min_non_null:
            return i
    return 0

async def read_file(file: UploadFile) -> dict:
    filename = file.filename.lower()
    file_bytes = await file.read()

    # FIX 2: Return cached result immediately if this exact file was seen before
    cache_key = _cache_key(file_bytes)
    if cache_key in _DOC_CACHE:
        logger.info("Cache hit for file '%s' (key=%s…)", filename, cache_key[:12])
        return _DOC_CACHE[cache_key]

    loop = asyncio.get_running_loop()

    def _parse():
        if filename.endswith(".csv") or filename.endswith(".xlsx"):
            if filename.endswith(".csv"):
                raw_df = pd.read_csv(StringIO(file_bytes.decode("utf-8")), header=None, comment="#")
            else:
                raw_df = pd.read_excel(BytesIO(file_bytes), header=None)

            header_row = detect_header_row(raw_df)
            df = raw_df.iloc[header_row + 1:].copy()
            df.columns = raw_df.iloc[header_row]
            df = df.dropna(how="all")
            df = df[df.notnull().sum(axis=1) > 1]
            df = df.replace([np.inf, -np.inf], None)
            df = df.where(pd.notnull(df), None)
            df = df.drop_duplicates()

            return {
                "filename": filename,
                "header_row_index": header_row,
                "rows": len(df),
                "columns": list(map(str, df.columns)),
                "data": df.to_dict(orient="records"),
            }

        elif filename.endswith(".docx") or filename.endswith(".doc"):
            try:
                document = Document(BytesIO(file_bytes))
            except Exception as e:
                raise ValueError("Could not parse Word Document.") from e

            seen = set()
            unique_lines = []

            for para in document.paragraphs:
                cleaned_text = para.text.strip()
                if cleaned_text and cleaned_text.lower() not in seen:
                    seen.add(cleaned_text.lower())
                    unique_lines.append(cleaned_text)

            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text and row_text.lower() not in seen:
                        seen.add(row_text.lower())
                        unique_lines.append(row_text)

            data_list = [{"document_line": line} for line in unique_lines]
            return {
                "filename": filename,
                "header_row_index": 0,
                "rows": len(data_list),
                "columns": ["document_line"],
                "data": data_list,
            }
        else:
            raise ValueError("Unsupported format.")

    result = await loop.run_in_executor(None, _parse)

    # FIX 2: Store parsed result in cache before returning
    _DOC_CACHE[cache_key] = result
    logger.info("Cache stored for file '%s' (key=%s…)", filename, cache_key[:12])
    return result


# ============================================================
# LLM helpers
# ============================================================

@with_retry  # FIX 3
async def identify_doc_type(format: dict) -> dict:
    prompt_doc = """
You will be given USER DATA. Identify if it is a task tracker or not.
If in doubt, consider it a success.

Respond with ONLY a JSON object — no markdown, no backticks.

Success: {{"status": 0, "message": "success"}}
Failure: {{"status": 1, "message": "failure"}}

USER DATA: {user_data}
"""
    chain = ChatPromptTemplate.from_template(prompt_doc) | _make_haiku()  # FIX 4
    try:
        raw = await chain.ainvoke({"user_data": format})
        return _parser.parse(raw.content)
    except Exception as e:
        return {"status": 2, "message": str(e)}


@with_retry  # FIX 3
async def club_tasks(format: dict) -> dict:
    preset_list = await get_preset_list()

    prompt_doc = """
Assign each task from TASK_LIST to the most relevant project category from PROJECT_LIST.
Use the scope field as a guide. Only use 'Miscellaneous' if no match exists.

RESPONSE FORMAT:
{{"exact task name": "exact project name", ...}}

PROJECT_LIST: {project_list}
TASK_LIST: {task_list}
"""
    chain = ChatPromptTemplate.from_template(prompt_doc) | _make_haiku()  # FIX 4
    raw = await chain.ainvoke({"task_list": format, "project_list": preset_list})
    parsed = _parser.parse(raw.content)
    project_tasks = {}
    for task, category in parsed.items():
        project_tasks.setdefault(category, []).append(task)

    meta = raw.response_metadata["usage"]
    inp, out = meta["input_tokens"], meta["output_tokens"]
    return {
        "tasks": project_tasks,
        "meta": {
            "model": "claude-haiku-4-5",
            "input_cost_usd": inp * (0.80 / 1_000_000),
            "output_cost_usd": out * (4.00 / 1_000_000),
            "total_cost_usd": (inp * 0.80 + out * 4.00) / 1_000_000,
        },
    }


# ============================================================
# PHASE 1: Chunk processor (Strict Extraction Only)
# ============================================================

def _chunk_data(data: list, chunk_size: int) -> list:
    return [data[i:i + chunk_size] for i in range(0, len(data), chunk_size)]


@with_retry  # FIX 3
async def _process_chunk(
    chunk: list,
    chunk_index: int,
    project_name: str,
    project_description: str,
    user_data: str,
    preset_list: list,
    is_first_chunk: bool,
) -> tuple:

    prompt_doc_chunk = """
    You are a data-extraction assistant processing a TASK TRACKER CHUNK.

    **YOUR ONLY JOB**:
    1. Extract the actionable tasks exactly as they appear in the tracker.
    2. Match every task to the most relevant preset from PRESET_LIST.
    3. Assign the most suitable agent from AGENT_LIST.
    4. Write clear, professional task descriptions based on the provided row.
    5. {project_name_instruction}

    **STRICT EXTRACTION RULES**:
    - DO NOT invent or add any missing tasks. Only extract what is present in the chunk.
    - Set origin to 1 for all extracted tasks.
    - Skip section headers, empty rows, and non-task content.
    - For deadline: extract the raw date string as-is.
    - For budget: extract the raw value including currency symbol. Only include if an actual monetary value is present in the row.

    **TASK NAME RULES** (most important):
    - NEVER use coded names like "Task_2_Marketing" or "Task_16_Booth"
    - Write the task name as a clear, action-oriented title.
    - Format: [Verb] + [What] + [Context if needed]
    - Examples:
        * "Design Social Media Banners for Launch Week"
        * "Finalize Booth Layout and Vendor Coordination"
        * "Book Travel for Campaign Event – Philadelphia"

    **RESPONSE FORMAT**:
    {{
        "project_name": "{project_name_value}",
        "preset_name": "Preset name from PRESET_LIST",
        "tasks": [
            {{
                "name": "Action-oriented descriptive task name",
                "description": "Clear professional description of what needs to be done",
                "origin": 1,
                "agent": {{"name": "agent name from AGENT_LIST", "id": agent_id_integer}},
                "details": {{
                    "assigned_to": "Person Name or null",
                    "deadline": "raw date string or null",
                    "status": "raw status string or null",
                    "budget": "raw budget value with currency or null"
                }}
            }}
        ]
    }}

    **DATA**:
    TASK TRACKER CHUNK: {chunk}
    USER DATA: {user_data}
    PRESET_LIST: {preset_list}
    AGENT_LIST: {agent_list}
    ADDITIONAL PROJECT DATA: {project_description}
    REFERENCE PROJECT NAME: {project_name}

    Agent assignment guide:
    - Design, mailer, social media, email, branding, website -> marketing (id: 4)
    - Order forms, templates, reports, shipping docs, planning -> business plan (id: 2)
    - Pricing, discounts, margins, costs, financial calculations -> financial (id: 3)
    - Sourcing, procurement, physical items, logistics -> procurement (id: 6)
    - Contracts, insurance, compliance, legal documents -> legal (id: 5)
    - Competitor research, attendee lists, market analysis -> market research (id: 1)
    """

    project_name_instruction = (
        "Create an extremely creative, catchy, and memorable project name that captures the essence of the work"
        if is_first_chunk
        else f"Use the project name: {project_name}"
    )
    project_name_value = "generate_creative_name" if is_first_chunk else project_name

    chain = ChatPromptTemplate.from_template(prompt_doc_chunk) | _make_haiku()  # FIX 4

    raw = await chain.ainvoke({
        "chunk": chunk,
        "user_data": user_data,
        "preset_list": preset_list,
        "agent_list": agents,
        "project_description": project_description,
        "project_name": project_name,
        "project_name_instruction": project_name_instruction,
        "project_name_value": project_name_value,
    })

    parsed = _parser.parse(raw.content)
    meta = raw.response_metadata["usage"]
    return parsed, meta["input_tokens"], meta["output_tokens"]


# ============================================================
# PHASE 3: Global Semantic Expansion
# ============================================================

@with_retry  # FIX 3
async def _refine_and_expand_tasks(
    existing_tasks: list,
    project_name: str,
    project_description: str,
    user_data: str,
    preset_list: list,
) -> tuple:

    simplified_tasks = [{"name": t["name"], "description": t["description"], "details": t.get("details", {})} for t in existing_tasks]

    prompt_expand = """
    You are an expert Project Manager. You are reviewing a raw list of tasks extracted from a document for the project: "{project_name}".
    Because these tasks were extracted in chunks, there are many duplicates, near-duplicates, and highly overlapping tasks in this list.
    
    **YOUR JOB**:
    1. **Deduplicate and Merge**: Go through all the existing tasks. Combine, merge, and deduplicate any tasks that share the same semantic goal, operational stage, or conceptually overlap (e.g., "Marketing Content" and "Design Assets" should merge). Think holistically. Set origin type to `user_upload` for these merged/verified tasks.
    2. **Expand**: Identify CRITICAL missing phases, prerequisites, or follow-up tasks that are completely absent from the existing list but are absolutely required for project success. Add ONLY 3 to 7 of these high-value tasks. Set origin type to `ai_generated` for these.
    
    **RULES**:
    - DO NOT drop any unique, distinct tasks from the existing list. Only squash tasks if they clearly overlap operationally.
    - Write **EXTREMELY CONCISE** descriptions. Do not be overly verbose. Limit descriptions to 10-15 words maximum. Remove all repetitive preamble words (like "Manage and execute").
    - Assign the most suitable agent from AGENT_LIST for every task.
    - Combine deadlines and assignees logically if you merge tasks. If any task has a budget, sum them up or keep the most logical aggregated budget.
    
    **RESPONSE FORMAT**:
    {{
        "tasks": [
            {{
                "name": "Action-oriented descriptive task name",
                "description": "Clear professional description of this task",
                "origin": {{"type": "user_upload" | "ai_generated", "confidence": 1.0 (for user_upload) or 0.8 (for ai_generated)}},
                "agent": {{"name": "agent name from AGENT_LIST", "id": agent_id_integer}},
                "details": {{
                    "assigned_to": ["Person 1", "Person 2"] or null,
                    "deadline": "date string" or null,
                    "status": "planned" or "unknown",
                    "budget": "budget string" or null
                }}
            }}
        ]
    }}

    **DATA**:
    EXISTING PROJECT TASKS: {existing_tasks}
    PROJECT DESCRIPTION: {project_description}
    USER DATA: {user_data}
    PRESET_LIST: {preset_list}
    AGENT_LIST: {agent_list}
    """

    chain = ChatPromptTemplate.from_template(prompt_expand) | _make_sonnet()  # FIX 4

    raw = await chain.ainvoke({
        "project_name": project_name,
        "existing_tasks": simplified_tasks,
        "project_description": project_description,
        "user_data": user_data,
        "preset_list": preset_list,
        "agent_list": agents,
    })

    parsed = _parser.parse(raw.content)
    meta = raw.response_metadata["usage"]
    return parsed.get("tasks", []), meta["input_tokens"], meta["output_tokens"]


# ============================================================
# Main orchestrator (The 3-Phase Pipeline)
# ============================================================

async def generate_doc_project(
    format: dict,
    project_name: str,
    project_description: str,
    user_data: str,
    doc_type: dict,
) -> dict:

    if doc_type["status"] != 0:
        yield {"error": "Failed in identification"}
        return

    preset_list = await get_preset_list()

    data = format.get("data", [])
    chunks = _chunk_data(data, CHUNK_SIZE)

    if not chunks:
        validated = ProjectModel(project_name=project_name, preset_name="", tasks=[])
        yield {
            "projects": [validated.model_dump()],
            "meta": {
                "model": "claude-haiku-4-5",
                "chunks_processed": 0,
                "total_tasks": 0,
                "input_cost_usd": 0.0,
                "output_cost_usd": 0.0,
                "total_cost_usd": 0.0,
            },
        }
        return

    try:
        yield {"status_message": "Starting parallel data extraction", "progress": 10}
        
        # --------------------------------------------------
        # PHASE 1: Full Parallel Chunk Extraction
        # --------------------------------------------------
        chunk_results = await asyncio.gather(*[
            _process_chunk(
                chunk=chunk, chunk_index=i,
                project_name=project_name,
                project_description=project_description,
                user_data=user_data,
                preset_list=preset_list,
                is_first_chunk=(i == 0),
            )
            for i, chunk in enumerate(chunks)
        ])

        extracted_tasks = []
        total_input = 0
        total_output = 0

        final_project_name = project_name
        final_preset_name = ""

        for i, (parsed, inp, out) in enumerate(chunk_results):
            total_input += inp
            total_output += out
            extracted_tasks.extend(parsed.get("tasks", []))
            if i == 0:
                final_project_name = parsed.get("project_name", project_name)
                final_preset_name = parsed.get("preset_name", "")

        yield {"status_message": f"Extracted {len(extracted_tasks)} initial tasks. Normalizing...", "progress": 50}

        # --------------------------------------------------
        # PHASE 2: Baseline Normalization & Deduplication
        # --------------------------------------------------
        baseline_tasks = _normalize_tasks(extracted_tasks)

        # --------------------------------------------------
        # PHASE 3: Global Semantic Refinement & Expansion
        # --------------------------------------------------
        yield {"status_message": "Merging duplicates and expanding missing tasks via LLM...", "progress": 70}

        fully_refined_tasks, exp_in, exp_out = await _refine_and_expand_tasks(
            existing_tasks=baseline_tasks,
            project_name=final_project_name,
            project_description=project_description,
            user_data=user_data,
            preset_list=preset_list,
        )
        total_input += exp_in
        total_output += exp_out

        # The LLM now returns the full list (merged user tasks + new AI tasks).
        # We can normalize it and deduplicate one final time just in case.
        normalized_all_tasks = _normalize_tasks(fully_refined_tasks)
        
        yield {"status_message": "Final programmatic deduplication and formatting...", "progress": 90}

        final_task_list = deduplicate_tasks(normalized_all_tasks)

        # --------------------------------------------------
        # FINAL VALIDATION
        # --------------------------------------------------
        validated = ProjectModel(
            project_name=final_project_name,
            preset_name=final_preset_name,
            tasks=final_task_list,
        )

        input_cost = total_input * (0.80 / 1_000_000)
        output_cost = total_output * (4.00 / 1_000_000)

        yield {
            "projects": [validated.model_dump()],
            "meta": {
                "model": "claude-haiku-4-5",
                "chunks_processed": len(chunks),
                "total_tasks": len(final_task_list),
                "input_cost_usd": input_cost,
                "output_cost_usd": output_cost,
                "total_cost_usd": input_cost + output_cost,
            },
        }

    except Exception as e:
        yield {"error": f"generate_doc_project failed: {e}"}